"""Parse CV via line metadata + LLM augmentation (notebook flow)."""

import json
import re
from pathlib import Path

from cv_parser.extract_text import extract_lines
from cv_parser.schemas import (
    CVParseResult,
    Presentation,
    PresentationRole,
    PresentationType,
    ProfessorMetadata,
    Publication,
    PublicationRole,
    PublicationStatus,
    PublicationType,
    Recognition,
)

MIME = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".doc": "application/msword",
}

AUGMENT_PROMPT = """AUGMENT the line metadata below. Each object has line_number, line, is_header_candidate, ml_category, nearest_ml_category.

CRITICAL: Do NOT drop any lines. Output array MUST have exactly {n} objects, same order, same line_number values.

1) CATEGORIZATION: Correct ml_category or nearest_ml_category if wrong.

2) MERGE LINES (only when needed): Only merge when lines were incorrectly broken up (e.g. author lists, journal names split). When merging: update "line" on first line; for continuation lines add "merged_into": line_number and "proposed": null. Keep all objects.

3) PROPOSED OBJECT: For each publication, presentation, or recognition item add "proposed" with the pydantic-shaped object. Use null for headers, pii, continuation lines.

Proposed shapes:
- publication: {{"year": 2020, "type": "book|journal|other", "status": "in_progress|published", "institution": "", "title": "", "role": "sole_author|co_author"}}
- presentation: {{"title": "", "year": 2020, "type": "conference|keynote|media|workshop|other", "role": "sole_presenter|co_presenter", "institution": ""}}
- recognition: {{"year": 2020, "title": "", "institution": ""}}

Rules: institution = publisher for books, journal name for articles. role sole_author when CV owner only author, co_author when multiple or "with X". Return ONLY the JSON array.

---
LINE METADATA:
{metadata_json}"""


def extract_line_metadata(document: bytes, mime: str, lines: list[str]) -> list[dict]:
    """Header hints (font size + bold). Header candidates categorized by keywords."""
    from io import BytesIO

    numbered = [(i, ln.strip()) for i, ln in enumerate(lines, 1) if ln.strip()]
    if not numbered:
        return []

    def _is_bold_font(fn):
        return fn and ("bold" in str(fn).lower() or "-b" in str(fn).lower() or "_b" in str(fn).lower())

    def _pdf_format_hints(doc, lns):
        try:
            import pdfplumber
        except ImportError:
            return {}
        plumber_lines, sizes = [], []
        with pdfplumber.open(BytesIO(doc)) as pdf:
            for page in pdf.pages:
                words = page.extract_words(extra_attrs=["size", "fontname"])
                if not words:
                    continue
                rows = {}
                for w in words:
                    top = int(w.get("top", 0) or 0)
                    rows.setdefault(top, []).append(w)
                for top in sorted(rows):
                    ws = rows[top]
                    ws_sorted = sorted(ws, key=lambda w: (w.get("x0", 0), w.get("x1", 0)))
                    text = " ".join(w.get("text", "") or "" for w in ws_sorted).strip()
                    if not text:
                        continue
                    szs = [float(w.get("size", 0) or 0) for w in ws if w.get("size") is not None]
                    max_sz = max(szs) if szs else 0
                    all_bold = bool(ws_sorted and all(_is_bold_font(w.get("fontname")) for w in ws_sorted))
                    if max_sz > 0:
                        sizes.append(max_sz)
                    plumber_lines.append((text, max_sz, all_bold))
        if not sizes:
            return {}
        thresh = sorted(sizes)[len(sizes) // 2] * 1.1
        used = [False] * len(plumber_lines)
        numd = [(i, ln.strip()) for i, ln in enumerate(lns, 1) if ln.strip()]

        def norm(s):
            return " ".join(s.split())

        hints = {}
        for num, ln in numd:
            ln_n = norm(ln)
            for j, row in enumerate(plumber_lines):
                if not used[j] and norm(row[0]) == ln_n:
                    sz, bold = row[1], row[2]
                    hints[num] = (sz >= thresh and sz > 0) or bold
                    used[j] = True
                    break
            else:
                hints[num] = False
        return hints

    def _docx_format_hints(doc, lns):
        from docx import Document

        doc = Document(BytesIO(doc))
        hints, n = {}, 1
        for p in doc.paragraphs:
            style = (p.style and p.style.name or "").lower()
            runs = p.runs
            all_bold = bool(runs and all(getattr(r.font, "bold", None) for r in runs))
            hints[n] = "heading" in style or (all_bold and len(p.text.strip()) < 80)
            n += 1
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        style = (p.style and p.style.name or "").lower()
                        runs = p.runs
                        all_bold = bool(runs and all(getattr(r.font, "bold", None) for r in runs))
                        hints[n] = "heading" in style or (all_bold and len(p.text.strip()) < 80)
                        n += 1
        return hints

    def _stem_line(line):
        from nltk.stem import SnowballStemmer

        stemmer = SnowballStemmer("english")
        words = re.findall(r"\b\w+\b", line.lower())
        return set(stemmer.stem(w) for w in words)

    def _categorize_header(line, kw_stems):
        line_stems = _stem_line(line)
        best_cat, best_score = "other", 0
        for cat in ("publication", "presentation", "recognition"):
            score = len(line_stems & kw_stems[cat])
            if score > best_score:
                best_score, best_cat = score, cat
        return best_cat if best_score > 0 else "other"

    from nltk.stem import SnowballStemmer

    stemmer = SnowballStemmer("english")
    kw = {
        "publication": {"publication", "publications", "journal", "paper", "papers", "published", "article", "articles", "book", "books", "chapter", "chapters", "review", "working", "submitted", "accepted", "jae", "jar", "jf", "accounting", "round"},
        "presentation": {"conference", "conferences", "workshop", "workshops", "keynote", "keynotes", "talk", "talks", "presentation", "presentations", "panel", "panels", "media", "interview", "interviews", "consortium", "faculty", "invited"},
        "recognition": {"mentions", "mentioned", "recognition", "recognitions", "award", "awards", "honor", "honors", "fellowship", "fellowships", "grant", "grants", "prize", "prizes", "fellow", "professor", "chair", "selected", "distinguished", "outstanding", "accomplishment", "accomplishments"},
    }
    kw_stems = {cat: {stemmer.stem(w) for w in words} for cat, words in kw.items()}

    format_hints = _pdf_format_hints(document, lines) if mime == "application/pdf" else {}
    if mime in ("application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"):
        format_hints = _docx_format_hints(document, lines)

    out = []
    for num, ln in numbered:
        fmt = format_hints.get(num, False)
        ml_cat = _categorize_header(ln, kw_stems)
        out.append({
            "line_number": num,
            "line": ln,
            "is_header_candidate": fmt,
            "ml_category": ml_cat if fmt else "",
        })
    out = sorted(out, key=lambda x: x["line_number"])
    headers = [(m["line_number"], m["ml_category"]) for m in out if m["is_header_candidate"] and m["ml_category"]]
    for m in out:
        if not m["is_header_candidate"]:
            n = m["line_number"]
            prev = [(h[0], h[1]) for h in headers if h[0] < n]
            nearest = sorted(prev, key=lambda h: n - h[0])[:1]
            m["nearest_ml_category"] = nearest[0][1] if nearest else ""
    return out


def _clean_metadata(line_metadata: list[dict]) -> list[dict]:
    """Drop 'other' header/non-header candidates; remove source, stemmed_words."""
    cleaned = []
    for m in line_metadata:
        if m.get("is_header_candidate") and m.get("ml_category") == "other":
            continue
        if not m.get("is_header_candidate") and m.get("nearest_ml_category") == "other":
            continue
        o = {k: v for k, v in m.items() if k not in ("source", "stemmed_words")}
        cleaned.append(o)
    return cleaned


def _parse_augmented_json(raw: str) -> list[dict]:
    t = raw.strip()
    if not t:
        raise ValueError("Empty LLM response")
    if "```json" in t:
        m = re.search(r"```json\s*([\s\S]*)", t)
        t = (m.group(1) or "").strip()
        if t.endswith("```"):
            t = t[:-3].strip()
    elif "```" in t:
        m = re.search(r"```\s*([\s\S]*?)```", t, re.DOTALL)
        t = m.group(1).strip() if m else t
    if t and not t.lstrip().startswith("["):
        m = re.search(r"\[[\s\S]*\]", t)
        t = m.group(0) if m else t

    def _repair(s):
        idx = s.rfind("},")
        if idx < 0:
            idx = s.rfind("}")
        if idx >= 0:
            try:
                fixed = s[: idx + 1].rstrip().rstrip(",") + "]"
                return json.loads(fixed)
            except json.JSONDecodeError:
                pass
        return None

    try:
        return json.loads(t)
    except json.JSONDecodeError as e:
        if "Unterminated" in str(e):
            repaired = _repair(t)
            if repaired is not None:
                return repaired
        raise


def _augmented_to_cv_result(augmented: list[dict], filename: str) -> CVParseResult:
    """Convert augmented line metadata (items with proposed) to CVParseResult."""
    pubs, pres, recs = [], [], []
    meta = ProfessorMetadata(filename=filename)
    for m in augmented:
        p = m.get("proposed")
        if not p:
            continue
        cat = m.get("ml_category") or m.get("nearest_ml_category") or "publication"
        try:
            if cat == "publication" and "status" in p:
                pubs.append(Publication(
                    year=int(p.get("year", 0) or 0),
                    type=PublicationType(p.get("type", "other")),
                    status=PublicationStatus(p.get("status", "published")),
                    institution=str(p.get("institution", "")),
                    title=str(p.get("title", "")),
                    role=PublicationRole(p.get("role", "co_author")),
                ))
            elif cat == "presentation" and "type" in p:
                pres.append(Presentation(
                    title=str(p.get("title", "")),
                    year=int(p.get("year", 0) or 0),
                    type=PresentationType(p.get("type", "other")),
                    role=PresentationRole(p.get("role", "co_presenter")),
                    institution=str(p.get("institution", "")),
                ))
            elif cat == "recognition" or ("title" in p and "institution" in p and "status" not in p and "type" not in p):
                recs.append(Recognition(
                    year=int(p.get("year", 0) or 0),
                    title=str(p.get("title", "")),
                    institution=str(p.get("institution", "")),
                ))
        except (ValueError, KeyError):
            pass
    return CVParseResult(metadata=meta, publications=pubs, presentations=pres, recognitions=recs)


def parse_cv_from_lines(
    path: Path,
    *,
    provider: str | None = None,
    model: str | None = None,
    api_key: str | None = None,
) -> CVParseResult:
    """Parse CV using line metadata + LLM augmentation. Returns CVParseResult with metadata.filename."""
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(path)
    suffix = path.suffix.lower()
    mime = MIME.get(suffix)
    if not mime:
        raise ValueError(f"Unsupported format: {suffix}. Use .pdf or .docx")
    document = path.read_bytes()

    lines = extract_lines(document, mime, merge_lines=False)
    line_metadata = extract_line_metadata(document, mime, lines)
    cleaned = _clean_metadata(line_metadata)
    if not cleaned:
        return CVParseResult(metadata=ProfessorMetadata(filename=path.name))

    metadata_json = json.dumps(cleaned, indent=2)
    prompt = AUGMENT_PROMPT.format(n=len(cleaned), metadata_json=metadata_json)

    from cv_parser.providers import get_provider

    prov = get_provider(provider=provider, model=model, api_key=api_key)
    raw, _ = prov.parse(
        document=metadata_json.encode("utf-8"),
        mime_type="text/plain",
        prompt=prompt,
        return_raw=True,
    )
    augmented = _parse_augmented_json(str(raw))
    if isinstance(augmented, dict) and "lines" in augmented:
        augmented = augmented["lines"]
    if not isinstance(augmented, list):
        raise RuntimeError(f"Expected JSON array, got {type(augmented)}")

    return _augmented_to_cv_result(augmented, path.name)
