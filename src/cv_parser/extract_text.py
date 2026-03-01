"""Extract every line from PDF or DOCX for raw output."""

import re
from pathlib import Path

_BULLET_RE = re.compile(r"^[•\-*]|^o\s{2,}|^\d+\.|^[a-z]\.", re.I)
_NUMERIC_SECTION_RE = re.compile(r"^\d+\.\s+")
_PAGE_HEADER_RE = re.compile(r".*Page\s+\d+.*", re.I)
_CONTINUATION_WORDS = frozenset({"of", "and", "in", "the", "or", "to", "for"})
_MAX_MERGE_LEN = 200


def _last_word(s: str) -> str:
    """Last word of line (lowercase) for continuation check."""
    parts = s.strip().rstrip(".,;:!?")
    return parts.split()[-1].lower() if parts else ""


def _is_continuation(line: str, next_stripped: str) -> bool:
    """Line signals continuation: next line should merge."""
    if not line or not next_stripped:
        return False
    r = line.rstrip()
    if r.endswith("-") and not r.endswith("--"):
        return True
    if r.endswith(",") or r.endswith("("):
        return True
    if "(" in r and ")" not in r:
        return True
    if _last_word(r) in _CONTINUATION_WORDS:
        return True
    if next_stripped[0].islower() or not next_stripped[0].isalnum():
        return True
    return False


def _is_section_start(stripped: str) -> bool:
    """Line starts new section: do not merge."""
    return bool(_BULLET_RE.match(stripped) or _NUMERIC_SECTION_RE.match(stripped))


def _is_page_header(stripped: str) -> bool:
    """Line looks like page header."""
    return bool(_PAGE_HEADER_RE.match(stripped))


def merge_wrapped_lines(lines: list[str]) -> list[str]:
    """Merge lines split by word wrap or hyphenation. Preserves blanks."""
    out = []
    i = 0
    while i < len(lines):
        cur = lines[i]
        if not cur.strip():
            out.append(cur)
            i += 1
            continue
        merged = cur
        j = i + 1
        while j < len(lines):
            nxt = lines[j]
            if not nxt.strip():
                nxt_stripped = lines[j + 1].strip() if j + 1 < len(lines) else ""
                if _is_continuation(merged.rstrip(), nxt_stripped) and nxt_stripped:
                    j += 1
                    continue
                break
            r = merged.rstrip()
            nxt_stripped = nxt.strip()
            if _is_section_start(nxt_stripped):
                break
            if nxt_stripped[0].islower() or not nxt_stripped[0].isalnum():
                merged = r + " " + nxt.lstrip()
            elif len(merged) > _MAX_MERGE_LEN:
                break
            elif r.endswith("-") and not r.endswith("--"):
                if _is_page_header(nxt_stripped):
                    j += 1
                    continue
                merged = r[:-1] + nxt.lstrip()
            elif (r and r[-1] in ".!?;") or not _is_continuation(r, nxt_stripped):
                break
            else:
                merged = r + " " + nxt.lstrip()
            j += 1
        out.append(merged)
        i = j
    return out


def extract_lines(document: bytes, mime: str, *, merge_lines: bool = True) -> list[str]:
    """Extract every line from document. Returns list of lines (preserves blanks for count)."""
    if mime == "application/pdf":
        lines = _pdf_lines(document)
    elif mime in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ):
        lines = _docx_lines(document)
    elif mime == "text/plain":
        lines = document.decode("utf-8").splitlines()
    else:
        raise ValueError(f"Unsupported mime: {mime}")
    if merge_lines and mime == "application/pdf":
        lines = merge_wrapped_lines(lines)
    return lines


def _pdf_lines(document: bytes) -> list[str]:
    from io import BytesIO

    try:
        import pdfplumber

        with pdfplumber.open(BytesIO(document)) as pdf:
            lines: list[str] = []
            for page in pdf.pages:
                text = page.extract_text(layout=True) or ""
                for line in text.splitlines():
                    lines.append(line)
            return lines
    except ImportError:
        pass

    from pypdf import PdfReader

    reader = PdfReader(BytesIO(document))
    lines = []
    for page in reader.pages:
        try:
            text = page.extract_text(extraction_mode="layout") or ""
        except (KeyError, Exception):
            text = page.extract_text() or ""
        for line in text.splitlines():
            lines.append(line)
    return lines


def _docx_lines(document: bytes) -> list[str]:
    from docx import Document
    from io import BytesIO

    doc = Document(BytesIO(document))
    lines: list[str] = []
    for p in doc.paragraphs:
        lines.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    lines.append(p.text)
    return lines


def write_raw_lines(document: bytes, mime: str, out_path: Path) -> None:
    """Extract all lines and write to file, one per line."""
    lines = extract_lines(document, mime)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text("\n".join(lines), encoding="utf-8")
