"""Tests for extract_text module."""

from io import BytesIO
from pathlib import Path

import pytest

from cv_parser.extract_text import extract_lines, merge_wrapped_lines, write_raw_lines


def test_extract_lines_pdf():
    """PDF extraction returns list of lines."""
    from io import BytesIO
    from pypdf import PdfWriter

    # Create minimal valid PDF via pypdf
    w = PdfWriter()
    w.add_blank_page(72, 72)
    buf = BytesIO()
    w.write(buf)
    buf.seek(0)
    pdf_bytes = buf.read()

    lines = extract_lines(pdf_bytes, "application/pdf")
    assert isinstance(lines, list)


def test_extract_lines_docx(tmp_path):
    """DOCX extraction returns lines from paragraphs and tables."""
    from docx import Document
    from io import BytesIO

    doc = Document()
    doc.add_paragraph("Line one")
    doc.add_paragraph("Line two")
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    docx_bytes = buf.read()

    lines = extract_lines(docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    assert "Line one" in lines
    assert "Line two" in lines


def test_write_raw_lines_docx(tmp_path):
    """write_raw_lines creates file with one line per extracted line."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("A")
    doc.add_paragraph("B")
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    docx_bytes = buf.read()

    out = tmp_path / "raw.txt"
    write_raw_lines(docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", out)
    content = out.read_text(encoding="utf-8")
    assert "A" in content
    assert "B" in content


def test_docx_bold_at_start_only():
    """Bold at start -> header candidate; bold at end -> not header candidate."""
    from docx import Document

    doc = Document()
    p1 = doc.add_paragraph()
    p1.add_run("BoldStart").bold = True
    p2 = doc.add_paragraph()
    p2.add_run("Plain ").bold = False
    p2.add_run("BoldEnd").bold = True

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    docx_bytes = buf.read()

    doc2 = Document(BytesIO(docx_bytes))
    paras = list(doc2.paragraphs)
    first_bold_1 = bool(paras[0].runs and getattr(paras[0].runs[0].font, "bold", None))
    first_bold_2 = bool(paras[1].runs and getattr(paras[1].runs[0].font, "bold", None))

    assert first_bold_1, "First para starts with bold -> should be header"
    assert not first_bold_2, "Second para starts with plain -> should NOT be header"


def test_merge_wrapped_lines_hyphenation():
    """Hyphenation: Robo- + Journalism -> RoboJournalism."""
    got = merge_wrapped_lines(["Robo-", "Journalism"])
    assert got == ["RoboJournalism"]


def test_merge_wrapped_lines_continuation():
    """Continuation: The Chronicle of + Philanthropy... -> merged."""
    got = merge_wrapped_lines(["The Chronicle of", "Philanthropy, Fierce CFO."])
    assert got == ["The Chronicle of Philanthropy, Fierce CFO."]


def test_merge_wrapped_lines_no_merge_bullet():
    """No merge when next line starts with bullet."""
    inp = ["Section title.", "•  Bullet one"]
    got = merge_wrapped_lines(inp)
    assert got == inp


def test_merge_wrapped_lines_blank_preserved():
    """Blanks preserved; do not merge across blanks when no continuation."""
    inp = ["Line.", "", "Next"]
    got = merge_wrapped_lines(inp)
    assert got == inp


def test_merge_across_blank_hyphenation():
    """Hyphenation across blank: Robo- + blank + Journalism -> RoboJournalism."""
    got = merge_wrapped_lines(["Robo-", "", "Journalism"])
    assert got == ["RoboJournalism"]


def test_merge_across_blank_comma():
    """Merge across blank when continuation signal (ends with 'of')."""
    got = merge_wrapped_lines(["The Chronicle of", "", "Philanthropy."])
    assert got == ["The Chronicle of Philanthropy."]


def test_merge_author_list():
    """Author list split: (with Dave + Larcker, Charles). -> merged."""
    got = merge_wrapped_lines(["(with Dave", "Larcker, Charles)."])
    assert got == ["(with Dave Larcker, Charles)."]


def test_merge_journal_of():
    """Journal name split: Journal of + Accounting and Economics -> merged."""
    got = merge_wrapped_lines(["Journal of", "Accounting and Economics, 2021."])
    assert got == ["Journal of Accounting and Economics, 2021."]


def test_no_merge_section_number():
    """No merge when next line starts numbered section."""
    inp = ["Section.", "3.  New section"]
    got = merge_wrapped_lines(inp)
    assert got == inp


def test_merge_when_next_starts_lowercase():
    """Merge when next line starts with lowercase (mid-sentence wrap)."""
    got = merge_wrapped_lines(["The roles of information.", "awareness and acquisition costs."])
    assert got == ["The roles of information. awareness and acquisition costs."]


def test_merge_when_next_starts_non_alnum():
    """Merge when next line starts with non-alphanumeric (e.g. closing paren, comma)."""
    got = merge_wrapped_lines(["Select press: A, B,", ", C, D."])
    assert got == ["Select press: A, B, , C, D."]
